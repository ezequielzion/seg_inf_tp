import zipfile
import io
import os
import shutil
import qrcode
import requests
import openpyxl
from openpyxl.writer.excel import ExcelWriter
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import CT_Inline, parse_xml, CT_Picture
from docx.shape import InlineShape
from docx.shared import Length
from docx.text.run import Run
from docx import Document
import subprocess
import os.path


def generador_exec(endpoint: str, userAgentData: str):
    is_using_windows = "Windows" in userAgentData
    # s = separador
    s = "\\" if is_using_windows else "/"
    with open(f"tmp{s}ejecutable-jugoso.py", "w" if os.path.isfile(f"tmp{s}ejecutable-jugoso.py") else "x") as f:
        f.write(f'import requests\nrequests.get("{endpoint}")')
    try:
        subprocess.run(["pyinstaller", f"tmp{s}ejecutable-jugoso.py", "--onefile"], shell=False)
        executable_path = f"dist{s}ejecutable-jugoso"
    except BaseException as e:
        print(e)
    
    res = None
    with open(f"dist{s}ejecutable-jugoso", mode='rb') as file:
        res = file.read()
    os.remove(f"tmp{s}ejecutable-jugoso.py")
    os.remove(f"dist{s}ejecutable-jugoso")
    os.remove("ejecutable-jugoso.spec")
    shutil.rmtree("build")
    return res

def add_linked_pic(r: Run, image_path: str) -> InlineShape:
    """
    Image will be inserted as character without embedding, just as link.
    :param r: run
    :param image_path:
        Seems like it has to be absolute path.as_uri() like "file:///full/path/file.jpg".
        It seems that it also works with relative path like "./folder/image.jpg"
    :return:
    """

    # create RELATION.
    relations = r.part.rels
    rel_id = relations._next_rId
    relations.add_relationship(reltype=RELATIONSHIP_TYPE.IMAGE, target=image_path, rId=rel_id, is_external=True)

    # Comment about pic_id from python-docx creators:
    # -- Word doesn't seem to use this, but does not omit it
    pic_id = 0

    # Next code taken from this method:
    # def new(cls, pic_id, filename, rId, cx, cy):
    # Just one line changed in order to replace `r:embed` with `r:link`.

    # The following lines were created to make variable names same as in python-docx method.
    filename = image_path  # Filename - something useless. will make it equal to image_path
    cx = 1
    cy = 1

    # Expand that code as CT_Picture.new(pic_id, filename, rId, cx, cy):
    pic = parse_xml(CT_Picture._pic_xml())
    pic.nvPicPr.cNvPr.id = pic_id
    pic.nvPicPr.cNvPr.name = filename

    # pic.blipFill.blip.embed = rId  # This line is replaced with next one
    pic.blipFill.blip.link = rel_id

    pic.spPr.cx = cx
    pic.spPr.cy = cy

    shape_id = r.part.next_id

    # Now from here: inline = cls.new(cx, cy, shape_id, pic)
    inline = CT_Inline.new(cx, cy, shape_id, pic)
    r._r.add_drawing(inline)

    # For embedding image there is get_or_add_image_part method.
    # We don't need it here.

    return InlineShape(inline)

def get_ngrok_url():
    try:
        res = requests.get("http://localhost:4040/api/tunnels")
        res_dict = res.json()
        return res_dict['tunnels'][0]['public_url']
    except:
        return None


def generador_excel(endpoint: str):
    wb_nuevo = openpyxl.Workbook()
    sheet = wb_nuevo.active
    sheet.cell(column=1, row=1, value=f'=webservice("{endpoint}")')
    excel_buffer = io.BytesIO()
    archive = zipfile.ZipFile(excel_buffer, 'w', zipfile.ZIP_DEFLATED, allowZip64=True)
    writer = ExcelWriter(wb_nuevo, archive)
    writer.save()
    return excel_buffer.getvalue()

def generador_word(endpoint: str):
    document = Document()
    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    run = p.add_run('bold')
    add_linked_pic(run, endpoint)
    word_buffer = io.BytesIO()
    document.save(word_buffer)
    return word_buffer.getvalue()

def generador_qr(endpoint: str):
    ngrok_url = get_ngrok_url()
    ngrok_url = ngrok_url if ngrok_url else endpoint
    exactas_url = 'https://campus.exactas.uba.ar/'
    
    url = f"{ngrok_url}/redirect?final_url={exactas_url}&endpoint={endpoint}"
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    png_buffer = io.BytesIO()
    img.save(png_buffer, 'PNG')
    png_buffer.seek(0)
    return png_buffer.getvalue()

def generador_epub(endpoint: str):
    # Create a new ZIP file in memory
    epub_buffer = io.BytesIO()
    with zipfile.ZipFile(epub_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
        # Set the metadata for the book
        mimetype = 'application/epub+zip'
        zipf.writestr("mimetype", mimetype)
        
        container = '''<?xml version="1.0"?>
        <container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
        <rootfiles>
            <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
        </rootfiles>
        </container>'''
        zipf.writestr("META-INF/container.xml", container)
        
        metadata = '''<?xml version="1.0"?>
        <package version="3.0" xml:lang="en" xmlns="http://www.idpf.org/2007/opf" unique-identifier="book-id">
        <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
            <dc:identifier id="book-id">urn:uuid:B9B412F2-CAAD-4A44-B91F-A375068478A0</dc:identifier>
            <meta refines="#book-id" property="identifier-type" scheme="xsd:string">uuid</meta>
            <meta property="dcterms:modified">2000-03-24T00:00:00Z</meta>
            <dc:language>en</dc:language>
            <dc:title>My Book</dc:title>
            <dc:creator>John Smith</dc:creator>
        </metadata>
        <manifest>
            <item id="text" href="text.xhtml" media-type="application/xhtml+xml"/>
            <item id="toc" href="../OEBPS/toc.ncx" media-type="application/x-dtbncx+xml"/>
        </manifest>
        <spine toc="toc">
            <itemref idref="text"/>
        </spine>
        </package>'''
        zipf.writestr("OEBPS/content.opf", metadata)
        
        toc = '''<?xml version="1.0"?>
        <ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">
        <head>
            <meta name="dtb:uid" content="urn:uuid:B9B412F2-CAAD-4A44-B91F-A375068478A0"/>
            <meta name="dtb:depth" content="1"/>
            <meta name="dtb:totalPageCount" content="0"/>
            <meta name="dtb:maxPageNumber" content="0"/>
        </head>
        <docTitle>
            <text>My Book</text>
        </docTitle>
        <navMap>
            <navPoint id="navpoint-1" playOrder="1">
            <navLabel>
                <text>Chapter 1</text>
            </navLabel>
            <content src="text.xhtml#xpointer(/html/body/section[1])"/>
            </navPoint>
            <navPoint id="navpoint-2" playOrder="2">
            <navLabel>
                <text>Chapter 2</text>
            </navLabel>
            <content src="text.xhtml#xpointer(/html/body/section[2])"/>
            </navPoint>
        </navMap>
        </ncx>'''
        zipf.writestr("OEBPS/toc.ncx", toc)
        
        text = f'''<?xml version="1.0" encoding="UTF-8" standalone="no"?>
        <!DOCTYPE html>
        <html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" xml:lang="en" lang="en">
        <head>
            <link href="{endpoint}" rel="stylesheet" crossorigin="anonymous">
            <title>My Book</title>
        </head>
        <body>
            <section><h1>Chapter 1</h1>
            <p>This is the text for chapter 1.</p>
            </section>
            <section><h1>Chapter 2</h1>
            <p>This is the text for chapter 2.</p>
            </section>
        </body>
        </html>'''
        zipf.writestr("OEBPS/text.xhtml", text)
    return epub_buffer.getvalue()

generadores = {
    "executable": generador_exec,
    "excel": generador_excel,
    "word": generador_word,
    "epub": generador_epub,
    "qr": generador_qr
}
